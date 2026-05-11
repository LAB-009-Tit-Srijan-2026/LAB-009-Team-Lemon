"""
Export summaries and conversations to external platforms
"""
from typing import Optional, Dict
import os
from notion_client import Client
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime
import markdown


class NotionExporter:
    """Export content to Notion"""
    
    def __init__(self):
        self.api_key = os.getenv("NOTION_API_KEY")
        self.client = Client(auth=self.api_key) if self.api_key else None
        self.database_id = os.getenv("NOTION_DATABASE_ID")
    
    async def export_summary(
        self, 
        video_title: str, 
        summary: str, 
        topics: list, 
        video_id: str
    ) -> Optional[str]:
        """
        Export summary to Notion database
        
        Args:
            video_title: Title of the video
            summary: Overall summary text
            topics: List of topic summaries
            video_id: Video identifier
        
        Returns:
            Notion page URL or None if failed
        """
        if not self.client or not self.database_id:
            return None
        
        try:
            # Create topics text
            topics_text = "\n".join([
                f"• {topic.get('title', 'Untitled')}: {topic.get('summary', '')}"
                for topic in topics
            ])
            
            # Create Notion page
            page = self.client.pages.create(
                parent={"database_id": self.database_id},
                properties={
                    "Name": {
                        "title": [
                            {
                                "text": {
                                    "content": video_title
                                }
                            }
                        ]
                    },
                    "Video ID": {
                        "rich_text": [
                            {
                                "text": {
                                    "content": video_id
                                }
                            }
                        ]
                    },
                    "Created": {
                        "date": {
                            "start": datetime.now().isoformat()
                        }
                    }
                },
                children=[
                    {
                        "object": "block",
                        "type": "heading_2",
                        "heading_2": {
                            "rich_text": [
                                {
                                    "type": "text",
                                    "text": {
                                        "content": "Summary"
                                    }
                                }
                            ]
                        }
                    },
                    {
                        "object": "block",
                        "type": "paragraph",
                        "paragraph": {
                            "rich_text": [
                                {
                                    "type": "text",
                                    "text": {
                                        "content": summary
                                    }
                                }
                            ]
                        }
                    },
                    {
                        "object": "block",
                        "type": "heading_2",
                        "heading_2": {
                            "rich_text": [
                                {
                                    "type": "text",
                                    "text": {
                                        "content": "Key Topics"
                                    }
                                }
                            ]
                        }
                    },
                    {
                        "object": "block",
                        "type": "paragraph",
                        "paragraph": {
                            "rich_text": [
                                {
                                    "type": "text",
                                    "text": {
                                        "content": topics_text
                                    }
                                }
                            ]
                        }
                    }
                ]
            )
            
            return page.get("url")
        except Exception as e:
            print(f"Error exporting to Notion: {e}")
            return None


class GoogleDocExporter:
    """Export content to Google Docs (requires Google Docs API setup)"""
    
    def __init__(self):
        # This would require Google OAuth setup
        # For now, we'll provide structure for future implementation
        self.service = None
    
    async def export_summary(
        self,
        video_title: str,
        summary: str,
        topics: list,
        video_id: str
    ) -> Optional[str]:
        """Export to Google Docs - requires Google OAuth setup"""
        # Implementation would go here with Google Docs API
        return None


class DocxExporter:
    """Export content to Word (.docx) files"""
    
    @staticmethod
    async def export_summary(
        video_title: str,
        summary: str,
        topics: list,
        video_id: str,
        conversation_history: Optional[list] = None
    ) -> bytes:
        """
        Export summary to Word document
        
        Returns:
            Document bytes
        """
        doc = Document()
        
        # Title
        title = doc.add_heading(video_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Video ID: ").bold = True
        meta.add_run(video_id)
        meta.add_run(f"\nGenerated: ").bold = True
        meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        
        # Summary
        doc.add_heading("Summary", 1)
        doc.add_paragraph(summary)
        
        # Topics
        doc.add_heading("Key Topics", 1)
        for topic in topics:
            doc.add_heading(topic.get("title", "Untitled"), 2)
            doc.add_paragraph(topic.get("summary", ""))
        
        # Conversation history if available
        if conversation_history:
            doc.add_heading("Q&A Session", 1)
            for exchange in conversation_history:
                doc.add_heading(f"Q: {exchange.get('question', '')}", 2)
                doc.add_paragraph(exchange.get("answer", ""))
        
        # Return document as bytes
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()


class MarkdownExporter:
    """Export content as Markdown"""
    
    @staticmethod
    async def export_summary(
        video_title: str,
        summary: str,
        topics: list,
        video_id: str,
        conversation_history: Optional[list] = None
    ) -> str:
        """Export summary as Markdown"""
        
        md = f"# {video_title}\n\n"
        md += f"**Video ID:** `{video_id}`  \n"
        md += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        # Summary
        md += "## Summary\n\n"
        md += f"{summary}\n\n"
        
        # Topics
        md += "## Key Topics\n\n"
        for topic in topics:
            md += f"### {topic.get('title', 'Untitled')}\n\n"
            md += f"{topic.get('summary', '')}\n\n"
        
        # Q&A
        if conversation_history:
            md += "## Q&A Session\n\n"
            for exchange in conversation_history:
                md += f"### Q: {exchange.get('question', '')}\n\n"
                md += f"{exchange.get('answer', '')}\n\n"
        
        return md


async def export_content(
    export_type: str,
    video_title: str,
    summary: str,
    topics: list,
    video_id: str,
    conversation_history: Optional[list] = None
) -> Optional[str | bytes]:
    """
    Export content to specified format
    
    Args:
        export_type: 'notion', 'docx', 'markdown', or 'google_docs'
        video_title: Title of the video
        summary: Overall summary
        topics: List of topic summaries
        video_id: Video identifier
        conversation_history: Optional Q&A history
    
    Returns:
        URL (for Notion), bytes (for DOCX), or string (for Markdown)
    """
    if export_type == "notion":
        exporter = NotionExporter()
        return await exporter.export_summary(video_title, summary, topics, video_id)
    
    elif export_type == "docx":
        return await DocxExporter.export_summary(
            video_title, summary, topics, video_id, conversation_history
        )
    
    elif export_type == "markdown":
        return await MarkdownExporter.export_summary(
            video_title, summary, topics, video_id, conversation_history
        )
    
    elif export_type == "google_docs":
        exporter = GoogleDocExporter()
        return await exporter.export_summary(video_title, summary, topics, video_id)
    
    return None
